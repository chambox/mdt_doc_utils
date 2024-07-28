import pandas as pd
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

class WordDocumentManager:
    def __init__(self, file_path=None):
        """
        Initializes the WordDocumentManager with a specified file path and loads the document if provided.

        Parameters
        ----------
        file_path : str, optional
            The path to the Word document to manage.

        Examples
        --------
        >>> manager = WordDocumentManager('path_to_your_document.docx')
        """
        if file_path:
            self.doc = Document(file_path)
        else:
            self.doc = Document()
        self.table_count = self._get_existing_table_count()

    def save(self, output_path):
        """
        Saves the modified document to a specified output path.

        Parameters
        ----------
        output_path : str
            The path where the modified document will be saved.

        Examples
        --------
        >>> manager.save('path_to_modified_document.docx')
        """
        self.doc.save(output_path)

    def _get_existing_table_count(self):
        """
        Counts the number of tables already in the document.

        Returns
        -------
        int
            The number of tables in the document.

        Examples
        --------
        >>> table_count = manager._get_existing_table_count()
        """
        return len(self.doc.tables)

    def clear_table_content_except(self, table_index, rows_to_keep):
        """
        Clears the content of all rows in the specified table except for the rows specified in rows_to_keep.

        Parameters
        ----------
        table_index : int
            The index of the table to modify.
        rows_to_keep : list of int
            A list of row indices whose content should not be cleared (0-based index).

        Examples
        --------
        >>> manager.clear_table_content_except(0, [0, 1])
        """
        table = self.doc.tables[table_index]
        for row_index, row in enumerate(table.rows):
            if row_index not in rows_to_keep:
                for cell in row.cells:
                    cell.text = ""

    def delete_empty_rows(self, table_index):
        """
        Deletes any rows in the specified table that are entirely empty.

        Parameters
        ----------
        table_index : int
            The index of the table to modify.

        Examples
        --------
        >>> manager.delete_empty_rows(0)
        """
        table = self.doc.tables[table_index]
        for row_index in range(len(table.rows) - 1, -1, -1):
            row = table.rows[row_index]
            if all(cell.text.strip() == "" for cell in row.cells):
                table._element.remove(row._element)

    def add_data_to_table(self, table_index, data_dict, start_row_index):
        """
        Adds data from a dictionary to the specified table starting from a given row index.

        Parameters
        ----------
        table_index : int
            The index of the table to modify.
        data_dict : dict
            A dictionary where keys are column headers and values are lists of column data.
        start_row_index : int
            The row index to start adding data (0-based index).

        Examples
        --------
        >>> data_to_add = {
        >>>     'Model Owner': ['M-O'],
        >>>     'Model Sponsor': ['M-S'],
        >>>     'Model Name': ['M-N']
        >>> }
        >>> manager.add_data_to_table(0, data_to_add, 1)
        """
        table = self.doc.tables[table_index]
        headers = list(data_dict.keys())
        num_rows = len(next(iter(data_dict.values())))

        # Ensure the table has enough rows
        for _ in range(num_rows - (len(table.rows) - start_row_index)):
            table.add_row()

        for i in range(num_rows):
            row = table.rows[start_row_index + i].cells
            for j, header in enumerate(headers):
                if j < len(row):
                    row[j].text = str(data_dict[header][i])

    def add_dataframe(self, dataframe: pd.DataFrame, caption=None):
        """
        Adds a pandas DataFrame to the document as a new table and applies RAG status coloring based on the 'Status' column.
        Also adds a caption with numbering to the table.

        Parameters
        ----------
        dataframe : pandas.DataFrame
            The pandas DataFrame to add to the document.
        caption : str, optional
            The caption text for the table.

        Examples
        --------
        >>> df = pd.DataFrame({
        >>>     'Task': ['Task 1', 'Task 2', 'Task 3'],
        >>>     'Status': ['Red', 'Amber', 'Green']
        >>> })
        >>> manager.add_dataframe(df, caption="Tasks and their RAG status")
        """
        # Add a caption before the table
        if caption:
            self.table_count += 1
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(f'Table {self.table_count}: {caption}')
            run.bold = True

        table = self.doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'
        # Add the header rows.
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(dataframe.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="B7DEE8"/>'.format(nsdecls('w'))))

        # Define RAG status colors
        status_colors = {
            'Red': 'FF0000',
            'Amber': 'FFA500',
            'Green': '00FF00'
        }

        # Add the data rows.
        for row in dataframe.itertuples(index=False):
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                if dataframe.columns[i].lower() == 'Status':  # Case insensitive match for 'status'
                    color = status_colors.get(value, 'FFFFFF')  # Default to white if status not found
                    row_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))

    def add_heading(self, text, level=1):
        """
        Adds a heading to the document.

        Parameters
        ----------
        text : str
            The text for the heading.
        level : int, optional
            The heading level (1 to 9).

        Examples
        --------
        >>> manager.add_heading("Document Heading", level=1)
        """
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text, bold=False, italic=False, underline=False, color=None, font_size=None):
        """
        Adds a paragraph to the document with optional formatting.

        Parameters
        ----------
        text : str
            The text for the paragraph.
        bold : bool, optional
            Set to True to make the text bold.
        italic : bool, optional
            Set to True to make the text italic.
        underline : bool, optional
            Set to True to underline the text.
        color : tuple, optional
            A tuple of (R, G, B) values to color the text.
        font_size : int, optional
            The font size of the text.

        Examples
        --------
        >>> manager.add_paragraph("This is a paragraph with some formatted text.", bold=True, italic=True, underline=True, color=(255, 0, 0), font_size=14)
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        if color:
            run.font.color.rgb = RGBColor(*color)
        if font_size:
            run.font.size = Pt(font_size)

    def add_bullet_points(self, items):
        """
        Adds a list of bullet points to the document.

        Parameters
        ----------
        items : list of str
            A list of strings, each representing a bullet point.

        Examples
        --------
        >>> manager.add_bullet_points(["Bullet point 1", "Bullet point 2", "Bullet point 3"])
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')

    def add_mixed_format_paragraph(self, parts):
        """
        Adds a paragraph with mixed formatting to the document.

        Parameters
        ----------
        parts : list of dict
            A list of dictionaries, each containing 'text' and optional formatting keys.

        Examples
        --------
        >>> mixed_format_parts = [
        >>>     {'text': 'This is ', 'bold': True},
        >>>     {'text': 'a paragraph ', 'italic': True, 'color': (0, 0, 255)},
        >>>     {'text': 'with mixed formatting.', 'underline': True, 'font_size': 12}
        >>> ]
        >>> manager.add_mixed_format_paragraph(mixed_format_parts)
        """
        paragraph = self.doc.add_paragraph()
        for part in parts:
            run = paragraph.add_run(part.get('text', ''))
            run.bold = part.get('bold', False)
            run.italic = part.get('italic', False)
            run.underline = part.get('underline', False)
            color = part.get('color', None)
            if color:
                run.font.color.rgb = RGBColor(*color)
            font_size = part.get('font_size', None)
            if font_size:
                run.font.size = Pt(font_size)
